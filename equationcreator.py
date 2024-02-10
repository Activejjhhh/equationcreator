import random
import operator
from docx import Document
from openpyxl import load_workbook
import pandas as pd

# Load your workbook
wb = load_workbook('C:/Users/Danie/Downloads/Copy of MVS Financial Model (Test v.3).xlsx' , data_only=True) #change this to the path of where your excel sheet is

# Select your sheet
ws = wb['Data Collection']

# Specify the column and the range of rows
column = 'I'
start_row = 4
end_row = 39  # Change this to the number of the last row you want to read



doc = Document()
variable_names = {
    1: 'Sector Variable Vague',
    2: 'Sector variable precise',
    3: 'Sector Variable Size',
    4: 'Market Variable Vague',
    5: 'Market variable precise',
    6: 'Market Variable Size',
    7: 'Economic Variable Vague',
    8: 'Economic variable precise',
    9: 'Economic Variable Size',
    10: 'This is a fail safe. If you are seeing this message, The code has gone wrong somewhere. Discard the equation and do not use it'
}

line_items = [
    "Cash & Equivalents",
    "Accounts Receivable",
    "Prepaid Expense",
    "PP&E",
    "Right Of Use Assets",
    "Goodwill",
    "Deferred Tax",
    "Accounts Payable",
    "Accrued Liabilities",
    "Taxes Payable",
    "Short Term Debt",
    "Long Term Debt",
    "Lease Liabilities",
    "Retained Earnings",
    "Net Sales/Revenue",
    "COGS",
    "R&D",
    "SG&A",
    "Interest Expense",
    "Nonoperating Income",
    "Tax Provision",
    "Net Income",
    "EPS",
    "Net Income",
    "Depreciation & Amortization",
    "Share-based Compensation",
    "Deferred Tax",
    "Income From Investments",
    "Accounts Receivable",
    "Inventory",
    "Prepaid Expense",
    "Accounts Payable",
    "Purchase of PP&E",
    "Proceeds from Asset Sales",
    "Issuance Of Debts",
    "Payments of Debts"
]


# Each individual in the population will be represented as a string of numbers and operations from your list, it also creates  an initial population of random individuals.
def generate_individual(numbers):
    length = random.randint(3, len(numbers))
    individual = []
    num_copy = [num for num in numbers if num != 0]  # Exclude zero
    for _ in range(length):
        if not num_copy:  # If num_copy is empty
            break  # Skip the current iteration
        num = random.choice(num_copy)
        num_copy.remove(num)
        individual.append(num)
        individual.append(random.choice(['+', '-', '*', '/']))
    return individual[:-1]



def generate_population(numbers, pop_size): 
    return [generate_individual(numbers) for _ in range(pop_size)]


#  #The fitness of an individual will be the inverse of the absolute difference between the evaluated expression and the target number.
def calculate_fitness(individual, target):
    expr = ''.join(map(str, individual))
    try:
        value = eval(expr)
        # Penalize individuals with duplicate numbers
        numbers = [individual[i] for i in range(0, len(individual), 2)]
        if len(numbers) != len(set(numbers)):
            return 0
        return 1 / abs(target - value)
    except (SyntaxError, ZeroDivisionError):
        return 0

# Select individuals for reproduction. Individuals with higher fitness should have a higher chance of being selected.
def select(population, fitnesses):
    return random.choices(population, weights=fitnesses, k=2)

# CCreate new individuals by combining parts of two selected individuals.
def crossover(parent1, parent2):
    index = random.randint(1, len(parent1) - 2)
    child1 = parent1[:index] + parent2[index:]
    child2 = parent2[:index] + parent1[index:]
    return child1, child2

# Randomly change parts of some individuals.
def mutate(individual, numbers):
    index = random.randrange(0, len(individual), 2)
    num_copy = numbers.copy()
    [num_copy.remove(num) for num in individual if num in num_copy]
    if num_copy:
        individual[index] = random.choice(num_copy)
    else:
        individual[index] = random.choice(numbers)


# Replace some individuals in the population with the newly created individuals
def replace(population, children):
    population.sort(key=lambda x: calculate_fitness(x, target))
    population[:2] = children



unprocessed_list = [[ws[f'{column}{row}'].value for column in ['I', 'J', 'L', 'N', 'O', 'Q', 'S', 'T', 'V', 'E']] for row in range(start_row, end_row+1)]

# Create list_of_lists where each number is multiplied by 100
list_of_lists = [[value * 100 if isinstance(value, (int, float)) else value for value in sublist] for sublist in unprocessed_list]

# Print the modified list of lists
for sublist in list_of_lists:
    print(sublist)


numbers = list_of_lists[0][:-1] 
target = list_of_lists[0][-1]

# Assuming 'numbers' is your list of variables
variable_mapping = {numbers[i]: i+1 for i in range(len(numbers))}

def standardize_equation(individual):
    standardized_individual = []
    for item in individual:
        if item in variable_mapping:
            # Replace the number with the corresponding variable name
            standardized_individual.append(variable_names[variable_mapping[item]])
        else:
            standardized_individual.append(item)
    return standardized_individual


# Genetic Algorithm
def genetic_algorithm(numbers, target, pop_size, generations):
    population = generate_population(numbers, pop_size)
    zero_fitness_warning_given = False  # Add this line
    for _ in range(generations):
        fitnesses = [calculate_fitness(individual, target) for individual in population]
        if sum(fitnesses) == 0:  # If all fitnesses are zero
            if not zero_fitness_warning_given:  # Only print the message once per sublist
                print("All individuals have zero fitness. Skipping...")
                zero_fitness_warning_given = True  # After printing the message once, don't print it again
            continue  # Skip to the next iteration
        parents = select(population, fitnesses)
        children = crossover(*parents)
        for child in children:
            if random.random() < 0.1:
                mutate(child, numbers)
        replace(population, children)
    best_individual = max(population, key=lambda x: calculate_fitness(x, target))
    standardized_best_individual = standardize_equation(best_individual)

    return ''.join(map(str, best_individual)), ''.join(map(str, standardized_best_individual))  # return both equations


for line_item, sublist in zip(line_items, list_of_lists):
    numbers_in_loop = sublist[:-1]  # All elements except the last one
    target_in_loop = sublist[-1]  # The last element
    variable_mapping = {numbers_in_loop[i]: i+1 for i in range(len(numbers_in_loop))}  # Update the variable mapping
    print(f"Processing {line_item}...")
    doc.add_paragraph(f"Processing {line_item}...")  # Add the line item to the Word document
    for _ in range(25): #change this for how many equations you want
        original, standardized = genetic_algorithm(numbers_in_loop, target_in_loop, pop_size=100, generations=1200) #change the generations to increase accuracy 
        standardized_equation = ''.join(standardized)  
        print(standardized_equation)
        doc.add_paragraph(standardized_equation)  
doc.save("equations.docx")

